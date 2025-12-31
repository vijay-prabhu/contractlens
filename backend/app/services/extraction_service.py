"""Text extraction service for PDF and DOCX documents."""
import io
from dataclasses import dataclass
from typing import Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument


@dataclass
class ExtractionResult:
    """Result of text extraction from a document."""
    text: str
    page_count: int
    metadata: dict


class ExtractionService:
    """Service for extracting text from various document formats."""

    def extract(self, content: bytes, file_type: str) -> ExtractionResult:
        """Extract text from document based on file type.

        Args:
            content: Raw bytes of the document
            file_type: Type of file ('pdf' or 'docx')

        Returns:
            ExtractionResult with extracted text and metadata

        Raises:
            ValueError: If file type is not supported
        """
        if file_type == "pdf":
            return self._extract_pdf(content)
        elif file_type == "docx":
            return self._extract_docx(content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _extract_pdf(self, content: bytes) -> ExtractionResult:
        """Extract text from PDF using PyMuPDF.

        Args:
            content: Raw PDF bytes

        Returns:
            ExtractionResult with text from all pages
        """
        doc = fitz.open(stream=content, filetype="pdf")

        pages_text = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            pages_text.append(text)

        metadata = {
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
            "subject": doc.metadata.get("subject", ""),
            "creator": doc.metadata.get("creator", ""),
            "producer": doc.metadata.get("producer", ""),
        }

        page_count = len(doc)
        doc.close()

        full_text = "\n\n".join(pages_text)

        return ExtractionResult(
            text=full_text.strip(),
            page_count=page_count,
            metadata=metadata,
        )

    def _extract_docx(self, content: bytes) -> ExtractionResult:
        """Extract text from DOCX using python-docx.

        Args:
            content: Raw DOCX bytes

        Returns:
            ExtractionResult with text from all paragraphs
        """
        doc = DocxDocument(io.BytesIO(content))

        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        # Extract metadata from core properties
        metadata = {}
        if doc.core_properties:
            metadata = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "subject": doc.core_properties.subject or "",
            }

        full_text = "\n\n".join(paragraphs)

        return ExtractionResult(
            text=full_text.strip(),
            page_count=1,  # DOCX doesn't have pages in the same sense
            metadata=metadata,
        )
